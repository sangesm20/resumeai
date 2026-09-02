import io
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

from huggingface_hub import InferenceClient
from core.config import settings


# =========================================================
# ROBUST EXTRACTION FOR MULTI-PAGE DOCX AND PDF RESUMES
# =========================================================

def extract_text_from_bytes(
    file_bytes: bytes,
    filename: str
) -> str:

    filename_lower = filename.lower()
    text_list = []

    if filename_lower.endswith(".pdf"):
        print("--- FORCING OCR FOR PDF ---")
        try:
            images = convert_from_bytes(
                file_bytes, 
                poppler_path=r'C:\Users\user\Downloads\Release-26.02.0-0\poppler-26.02.0\Library\bin'
            )
            for idx, img in enumerate(images):
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    text_list.append(ocr_text)
        except Exception as e:
            print(f"PDF OCR Error: {e}")

    elif filename_lower.endswith(".docx"):
        print("--- EXTRACTING FULL DOCX CONTENT (INCLUDING TABLES/COLUMNS) ---")
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # 1. Extract from all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_list.append(para.text)
            
            # 2. Extract from tables (since modern resume templates use tables/columns)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_list.append(cell.text)
                            
        except Exception as e:
            print(f"DOCX Extraction Error: {e}")

    elif filename_lower.endswith((".png", ".jpg", ".jpeg")):
        image = Image.open(io.BytesIO(file_bytes))
        ocr_text = pytesseract.image_to_string(image)
        if ocr_text.strip():
            text_list.append(ocr_text)

    else:
        raise ValueError("Unsupported file format")

    final_text = "\n".join(text_list).strip()
    print(f"--- FINAL EXTRACTED TEXT LENGTH: {len(final_text)} ---")
    return final_text


# =========================================================
# EMBEDDING GENERATION
# =========================================================

def generate_embedding(
    text: str
) -> list[float]:

    if not text or not text.strip():
        raise ValueError("Text cannot be empty")

    client = InferenceClient(
        provider="hf-inference",
        api_key=settings.HF_TOKEN
    )

    result = client.feature_extraction(
        text[:12000],
        model=settings.HF_MODEL
    )

    embedding = result.tolist()

    if (
        isinstance(embedding, list)
        and len(embedding) == 1
        and isinstance(embedding[0], list)
    ):
        embedding = embedding[0]

    return [
        float(value)
        for value in embedding
    ]